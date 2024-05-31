import camelot
from docx import Document

# Specify the path to the PDF file with tables
pdf_file = 'your_pdf_file.pdf'

# Create a Word document
doc = Document()

# Extract tables from the PDF file with adjusted settings
tables = camelot.read_pdf(pdf_file, pages='all', flavor='stream', layout_kwargs={'detect_vertical': True, 'all_texts': True})

# Iterate through each table
for table in tables:
    # Add a table to the Word document
    if not table.df.empty:
        table_df = table.df
        table_rows = [[str(cell) for cell in row] for row in table_df.values.tolist()]
        doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        table = doc.tables[-1]
        for i, row in enumerate(table_rows):
            row_cells = table.rows[i].cells
            for j, cell in enumerate(row):
                row_cells[j].text = cell

# Save the Word document
doc.save('your_pdf_file.docx')
